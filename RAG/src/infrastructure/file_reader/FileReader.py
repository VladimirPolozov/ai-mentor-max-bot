import os
from typing import List

from config import Config
from RAG.src.core.entities import Document
from RAG.src.core.interfaces import IFileReader
from docx import Document as DocxDocument


class FileReader(IFileReader):
    def read_files(self, directory_path: str) -> List[Document]:
        documents = []

        if not os.path.exists(directory_path):
            raise ValueError(f"Directory {directory_path} does not exist")

        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)

            if filename.endswith('.txt'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        content = file.read()
                except Exception as e:
                    print(f"Ошибка при чтении TXT файла {filename}: {e}")
                    continue

            elif filename.endswith('.docx'):
                try:
                    docx_file = DocxDocument(file_path)
                    content = "\n".join(p.text for p in docx_file.paragraphs if p.text.strip())
                except Exception as e:
                    print(f"Ошибка при чтении DOCX файла {filename}: {e}")
                    continue

            else:
                continue

            document = Document(
                content=content,
                metadata={
                    'file_name': filename,
                    'file_path': file_path,
                    'file_size': os.path.getsize(file_path)
                }
            )
            documents.append(document)

        return documents

if __name__ == "__main__":
    txt_file_reader = FileReader()
    config = Config()
    txt_file_reader.read_files(config.DATA_DIRECTORY)